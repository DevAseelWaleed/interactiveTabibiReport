# -*- coding: utf-8 -*-
import os, sys, json
sys.stdout.reconfigure(encoding='utf-8')

base_dir = os.path.dirname(os.path.abspath(__file__))

from docx import Document

docx_files = [f for f in os.listdir(base_dir) if f.endswith('.docx')]

all_data = {}

for docx_file in docx_files:
    full_path = os.path.join(base_dir, docx_file)
    doc = Document(full_path)
    
    file_data = {
        'name': docx_file,
        'paragraphs': [],
        'tables': []
    }
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            file_data['paragraphs'].append({
                'index': i,
                'style': para.style.name,
                'text': para.text
            })
    
    for t_idx, table in enumerate(doc.tables):
        table_data = {
            'index': t_idx,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'data': []
        }
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_data['data'].append(cells)
        file_data['tables'].append(table_data)
    
    all_data[docx_file] = file_data

output_path = os.path.join(base_dir, 'extracted_data.json')
with open(output_path, 'w', encoding='utf-8') as f:
    json.dump(all_data, f, ensure_ascii=False, indent=2)

print(f"Data exported to {output_path}")
print(f"Files: {list(all_data.keys())}")
for name, data in all_data.items():
    print(f"  {name}: {len(data['paragraphs'])} paragraphs, {len(data['tables'])} tables")
