# -*- coding: utf-8 -*-
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

base_dir = os.path.dirname(os.path.abspath(__file__))

# List files
print("=== FILES IN DIRECTORY ===")
for f in os.listdir(base_dir):
    full = os.path.join(base_dir, f)
    if os.path.isfile(full):
        size = os.path.getsize(full)
        print(f"  {f} ({size:,} bytes)")

# Find the docx file
docx_files = [f for f in os.listdir(base_dir) if f.endswith('.docx')]
print(f"\nDOCX files found: {docx_files}")

# Read the DOCX
from docx import Document
for docx_file in docx_files:
    full_path = os.path.join(base_dir, docx_file)
    print(f"\n\n{'='*60}")
    print(f"READING: {docx_file}")
    print('='*60)
    doc = Document(full_path)
    
    print("\n--- PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{i}: [{para.style.name}] {para.text}")
    
    print(f"\n--- TABLES ({len(doc.tables)} total) ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n  Table {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols):")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            print(f"    Row {r_idx}: {cells}")
