# -*- coding: utf-8 -*-
"""
Upgrade Word document generator, PPTX generators, and Web Slides to contain all details from Pages 30 to 47.
"""
import os, sys, shutil
from docx import Document
from docx.shared import Inches as DInches, Pt as DPt, RGBColor as DRGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')
base_dir = r"e:\Work\زبون تقرير نصف سنوي طبيبي"
v2_dir = os.path.join(base_dir, "التقرير_الاحترافي_المطور")
v1_dir = os.path.join(base_dir, "التقرير_الجديد")

PRIMARY = DRGBColor(107, 29, 58)       # #6B1D3A Burgundy
SECONDARY = DRGBColor(201, 169, 110)   # #C9A96E Gold
TEXT_DARK = DRGBColor(45, 45, 45)
GRAY_BG = "F4EFE6"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def add_rtl_p(doc, text, bold=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = DPt(space_after)
    p.paragraph_format.bidi = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = DPt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_rtl_heading(doc, text, level=1):
    colors = {1: PRIMARY, 2: SECONDARY, 3: PRIMARY}
    sizes = {1: 17, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = DPt(14)
    p.paragraph_format.space_after = DPt(6)
    p.paragraph_format.bidi = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = DPt(sizes.get(level, 13))
    run.font.bold = True
    run.font.color.rgb = colors.get(level, PRIMARY)
    return p

def build_comprehensive_word_doc(docx_path):
    print(f"Building Comprehensive Word Document: {docx_path}")
    doc = Document()

    # Configure Margins (0.8 in)
    for s in doc.sections:
        s.top_margin = DInches(0.8)
        s.bottom_margin = DInches(0.8)
        s.left_margin = DInches(0.8)
        s.right_margin = DInches(0.8)

    # 1. Cover
    p_title = add_rtl_p(doc, "جمعية طبيبي الأهلية بالمدينة المنورة", bold=True, color=SECONDARY, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_p(doc, "التقرير النصف سنوي الشامل والتدقيق الاستراتيجي لعام ٢٠٢٦م", bold=True, color=PRIMARY, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_rtl_p(doc, "«مدعم بنتائج التدقيق الأدائي المقارن ومطابقة مستهدفات الخطة الاستراتيجية والتشغيلية»", bold=True, color=SECONDARY, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_rtl_p(doc, "« ثـقـة  •  أثــر  •  اسـتـدامـة »", bold=True, color=PRIMARY, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_rtl_p(doc, "الفترة: من ١ يناير إلى ٣٠ يونيو ٢٠٢٦م | ترخيص المركز الوطني: (١٠٠٠٧٣٠٧٠٠)", color=TEXT_DARK, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_p(doc, "إعداد: أ. بيان بن سعد المحمدي - المدير التنفيذي | إشراف مجلس الإدارة", color=TEXT_DARK, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # 2. Executive Summary & Leadership
    add_rtl_heading(doc, "أولاً: القيادة الرشيدة ومجلس الإدارة (الرؤية والتمكين)", level=1)
    add_rtl_p(doc, "تسترشد جمعية طبيبي الأهلية بتوجيهات القيادة الرشيدة في تمكين القطاع غير الربحي ليكون شريكاً ريادياً في التنمية الصحية ورفع جودة الحياة بطيبة الطيبة.", color=TEXT_DARK, size=11)
    add_rtl_p(doc, "• خادم الحرمين الشريفين الملك سلمان بن عبدالعزيز: «ما يميز هذه البلاد هو حرص قادتها على الخير والتشجيع عليه، وما نراه من مؤسسات خيرية في مختلف المجالات إلا جانب من الجوانب المشرقة لبلادنا.»", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "• ولي العهد رئيس مجلس الوزراء الأمير محمد بن سلمان: «نهدف للوصول إلى قطاع غير ربحي مهم، مبادر وداعم ومؤثر في التعليم والصحة والثقافة والمجالات البحثية، وسنعتمد عليه بشكل رئيسي.»", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "• أمير منطقة المدينة المنورة الأمير سلمان بن سلطان: «نسعد بالإنجازات التي حققتها الجمعيات الأهلية على مستوى المنطقة باعتبارها شريكاً استراتيجياً في تحسين جودة الحياة.»", bold=True, color=PRIMARY, size=10.5)

    add_rtl_heading(doc, "ثانياً: نتائج التدقيق الاستراتيجي وبطاقة الأداء المتوازن (BSC)", level=1)
    add_rtl_p(doc, "أظهر التدقيق المقارن بين الخطة الاستراتيجية والتشغيلية والمنجز الفعلي للنصف الأول تحقيق نسبة إنجاز موزونة قدرها ٤١.٧٣٪ (بعد مواءمة المستهدف المالي السنوي إلى ١.٥٢٧ مليون ريال).", color=TEXT_DARK, size=11)

    # BSC Table
    t_bsc = doc.add_table(rows=6, cols=4)
    t_bsc.alignment = WD_TABLE_ALIGNMENT.CENTER
    bsc_headers = ["منظور بطاقة الأداء المتوازن (BSC)", "الوزن النسبي", "نسبة إنجاز المنظور", "الدرجة المحققة من ١٠٠"]
    for j, h in enumerate(bsc_headers):
        t_bsc.rows[0].cells[j].paragraphs[0].text = h
        set_cell_background(t_bsc.rows[0].cells[j], "6B1D3A")
        t_bsc.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = DRGBColor(255,255,255)
        t_bsc.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_bsc.rows[0].cells[j].paragraphs[0].paragraph_format.bidi = True
        t_bsc.rows[0].cells[j].paragraphs[0].runs[0].font.size = DPt(10)

    bsc_rows = [
        ("١. محور الأثر والبرامج الطبية", "٤٠٪", "١٣.٩٥٪", "٥.٥٨ من ٤٠"),
        ("٢. المحور المالي والموازنة التشغيلية", "٣٠٪", "٦٣.٠١٪", "١٨.٩٠ من ٣٠"),
        ("٣. محور العمليات والشراكات المؤسسية", "١٥٪", "٥٥.٠٠٪", "٨.٢٥ من ١٥"),
        ("٤. محور الحوكمة والقدرات المؤسسية", "١٥٪", "٦٠.٠٠٪", "٩.٠٠ من ١٥"),
        ("الإجمالي العام الموزون للأداء الاستراتيجي (H1 2026)", "١٠٠٪", "٤١.٧٣٪", "٤١.٧٣ من ١٠٠ (بعد مواءمة المستهدف)")
    ]
    for i, row in enumerate(bsc_rows):
        r = t_bsc.rows[i+1]
        for j, val in enumerate(row):
            r.cells[j].paragraphs[0].text = val
            r.cells[j].paragraphs[0].paragraph_format.bidi = True
            r.cells[j].paragraphs[0].runs[0].font.size = DPt(9.5)
            if i == 4:
                set_cell_background(r.cells[j], "E8DFC8")
                r.cells[j].paragraphs[0].runs[0].font.bold = True
            elif i % 2 == 1:
                set_cell_background(r.cells[j], "F8F6F0")
    doc.add_page_break()

    # 3. 14 Strategic Matrix Table
    add_rtl_heading(doc, "ثالثاً: مصفوفة مطابقة الخطة الاستراتيجية بالمنجز الفعلي (١٤ مؤشراً معتمداً)", level=1)
    matrix_headers = ["م", "الهدف / النشاط", "المؤشر المعتمد", "المستهدف السنوي", "مستهدف H1", "المنجز الفعلي", "نسبة الإنجاز", "الحالة"]
    t_mat = doc.add_table(rows=15, cols=8)
    t_mat.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(matrix_headers):
        t_mat.rows[0].cells[j].paragraphs[0].text = h
        set_cell_background(t_mat.rows[0].cells[j], "6B1D3A")
        t_mat.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = DRGBColor(255,255,255)
        t_mat.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_mat.rows[0].cells[j].paragraphs[0].paragraph_format.bidi = True
        t_mat.rows[0].cells[j].paragraphs[0].runs[0].font.size = DPt(9)
    
    matrix_rows = [
        ("١", "الإيرادات الكلية (المعدلة)", "إجمالي الدخل (ريال)", "١,٥٢٧,٠٠٠", "٧٦٣,٥٠٠", "٥٨٢,١٦٧.٥٢", "٧٦.٢٥٪", "متقدم وجيد"),
        ("٢", "إيرادات الموازنة التشغيلية", "إيراد الموازنة (ريال)", "١,٥٢٧,٠٠٠", "١,٥٢٧,٠٠٠", "٥٨٢,١٦٧.٥٢", "٣٨.١٢٪", "متأخر"),
        ("٣", "المساعدات العلاجية المباشرة", "مبالغ المساعدات (ريال)", "١,٥٠٠,٠٠٠", "٧٥٠,٠٠٠", "٢٠٨,٦٠٥.٠٠", "٢٧.٨١٪", "متأخر حرِج"),
        ("٤", "المستفيدون من الخدمات الصحية", "عدد المستفيدين (فرد)", "٣٦,٦٠٦", "١٨,٣٠٣", "٧ مستفيدين", "٠.٠٣٨٪", "متعثر تماماً"),
        ("٥", "الاستشارات الطبية والدوائية", "عدد الاستشارات", "١,٢٠٠", "٦٠٠", "٠ استشارة", "٠.٠٠٪", "لم يبدأ"),
        ("٦", "الدراسات واستطلاعات الرأي", "عدد الدراسات", "٦ دراسات", "٣ دراسات", "٠ دراسة", "٠.٠٠٪", "لم يبدأ"),
        ("٧", "ساعات وقيمة التطوع", "ساعات وقيمة التطوع", "٣,٠٠٠ س (٢٠٢ ألف)", "١,٥٠٠ س (١٠١ ألف)", "٤ فرص تطوعية", "غير مدققة", "متعثر"),
        ("٨", "تفعيل محفظة البرامج", "عدد البرامج النشطة", "١٠ برامج", "١٠ برامج", "برنامج واحد فقط", "١٠.٠٠٪", "متعثر"),
        ("٩", "عقد الشراكات الصحية", "عدد الشراكات", "٩ شراكات", "٩ شراكات", "٩ شراكات مفعلة", "١٠٠.٠٠٪", "مكتمل"),
        ("١٠", "توطين الوظائف والكادر", "نسبة التوطين (٪)", "١٠٠٪", "١٠٠٪", "١٠٠٪ (٣ موظفين)", "١٠٠.٠٠٪", "مكتمل"),
        ("١١", "تدريب وتأهيل الكادر", "عدد الدورات", "٤ دورات", "٢ دورة", "٨ دورات", "٤٠٠.٠٠٪", "متقدم ومكتمل"),
        ("١٢", "التحول الرقمي والمحاسبي", "تطبيق نظام سحابي", "نظام قيود", "نظام قيود", "تم تشغيل قيود", "١٠٠.٠٠٪", "مكتمل"),
        ("١٣", "معايير الحوكمة ومنصة نوى", "نسبة الامتثال", "١٠٠٪", "٥٠٪", "طلب استشاري", "٢٥.٠٠٪", "متأخر"),
        ("١٤", "تنويع مصادر الدخل الذاتي", "عدد مصادر الدخل", "٦ مصادر", "٦ مصادر", "٦ مصادر نشطة", "٦٦.٦٧٪", "قيد التنفيذ")
    ]
    for i, row in enumerate(matrix_rows):
        r = t_mat.rows[i+1]
        for j, val in enumerate(row):
            r.cells[j].paragraphs[0].text = val
            r.cells[j].paragraphs[0].paragraph_format.bidi = True
            r.cells[j].paragraphs[0].runs[0].font.size = DPt(8.5)
            if i % 2 == 1:
                set_cell_background(r.cells[j], "F8F6F0")
    doc.add_page_break()

    # 4. Human Resources & Vacancies (Pages 30-31)
    add_rtl_heading(doc, "رابعاً: الموارد البشرية والوظائف الشاغرة (ص ٣٠-٣١)", level=1)
    add_rtl_p(doc, "• الكادر الوظيفي الحالي (٣ موظفين رسميين + محاسب متعاون بنسبة توطين ١٠٠٪):", bold=True, color=PRIMARY, size=11)
    add_rtl_p(doc, "  ١. أ. بيان سعد المحمدي - المدير التنفيذي (موظف رسمي).\n  ٢. أ. غدير أحمد الحربي - المسؤول المالي والمشرفة على البرامج والمشاريع (موظفة رسمية).\n  ٣. أ. طراد محمد سمان - سكرتير تنفيذي (موظف رسمي).\n  ٤. أ. محمد الحسن بشير - محاسب قانوني متعاون.", color=TEXT_DARK, size=10.5)
    
    add_rtl_p(doc, "• خطة الشواغر والاحتياج الوظيفي (٣ وظائف أساسية):", bold=True, color=PRIMARY, size=11)
    t_vac = doc.add_table(rows=4, cols=3)
    t_vac.alignment = WD_TABLE_ALIGNMENT.CENTER
    vac_headers = ["المسمى الوظيفي", "الحالة الراهنة", "الإجراء والترشيح"]
    for j, h in enumerate(vac_headers):
        t_vac.rows[0].cells[j].paragraphs[0].text = h
        set_cell_background(t_vac.rows[0].cells[j], "6B1D3A")
        t_vac.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = DRGBColor(255,255,255)
        t_vac.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        t_vac.rows[0].cells[j].paragraphs[0].paragraph_format.bidi = True
        t_vac.rows[0].cells[j].paragraphs[0].runs[0].font.size = DPt(9.5)
    
    vac_data = [
        ("موظف/ـة علاقات عامة وإعلام", "قيد الترسيم", "تم ترشيح أ. فيصل الجهني (يعمل كمتعاون بدون أجر منذ ٣ أشهر لحين الترسيم)."),
        ("موظف/ـة تنمية موارد مالية", "شاغرة (تحتاج استقطاب)", "استقطاب كفاءة متخصصة لرفع قدرات كتابة المشاريع والمنح وجلب التمويل."),
        ("موظف/ـة موارد بشرية", "شاغرة (تحتاج استقطاب)", "وظيفة أساسية لتنظيم شؤون الموظفين، التدريب، ولوائح العمل والتطوع.")
    ]
    for i, row in enumerate(vac_data):
        r = t_vac.rows[i+1]
        for j, val in enumerate(row):
            r.cells[j].paragraphs[0].text = val
            r.cells[j].paragraphs[0].paragraph_format.bidi = True
            r.cells[j].paragraphs[0].runs[0].font.size = DPt(9)
            if i % 2 == 1:
                set_cell_background(r.cells[j], "F8F6F0")
    
    # 5. Administrative Achievements & HQ Relocation (Pages 32-33)
    add_rtl_heading(doc, "خامساً: الإنجازات الإدارية والمالية والانتقال للمقر الجديد (ص ٣٢-٣٣)", level=1)
    add_rtl_p(doc, "• الانتقال للمقر الجديد والوفر المالي: تم الانتقال إلى المقر الحالي بإيجار سنوي ٤٥,٠٠٠ ريال (بدعم نائب رئيس المجلس) مقارنة بـ ٧٠,٠٠٠ ريال سابقاً، محققاً وفراً سنوياً قدره ٢٥,٠٠٠ ريال.", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "• قائمة الإنجازات الإدارية والمالية الـ (١٢) المنجزة:", bold=True, color=PRIMARY, size=10.5)
    admin_list = [
        "١. أتمتة العمل المحاسبي وإدخال برنامج «قيود» المجاز من المركز الوطني.",
        "٢. بناء وتأسيس الشجرة المحاسبية وفق الدليل المعتمد للجمعيات الخيرية.",
        "٣. إقفال الأرباع السنوية لعام ٢٠٢٥م وتدقيقها محاسبياً بكفاءة عالية.",
        "٤. إعداد الموازنة التقديرية لعام ٢٠٢٦م واعتمادها من المجلس والعمومية.",
        "٥. تنفيذ جرد مالي وحصري شامل لكافة أصول الجمعية وتوثيقها رسمياً.",
        "٦. إعداد القوائم المالية واعتمادها من المركز الوطني للقطاع غير الربحي.",
        "٧. أرشفة جميع مستندات وملفات الجمعية ورقياً وإلكترونياً.",
        "٨. إنشاء وتفعيل نظام الاتصالات الإدارية الصادر والوارد لكافة المعاملات.",
        "٩. تطوير الموقع الإلكتروني بما يتوافق مع متطلبات الحوكمة والخصوصية.",
        "١٠. تطوير السياسات واللوائح الداخلية ونشرها بالبوابة الرسمية.",
        "١١. متابعة الالتزام والحوكمة (جاري العمل عليها لرفع التصنيف).",
        "١٢. إعادة هيكلة اللجان والاكتفاء بلجنتين: (التنفيذية، ولجنة المساعدات الطبية)."
    ]
    for item in admin_list:
        add_rtl_p(doc, item, color=TEXT_DARK, size=10)
    doc.add_page_break()

    # 6. Resource Development & Grants Pipeline (Pages 34-35)
    add_rtl_heading(doc, "سادساً: تنمية الموارد والمنح وفرص منصة إحسان (ص ٣٤-٣٥)", level=1)
    add_rtl_p(doc, "• إجمالي المنح المرفوعة: ٢٧ منحة | المنح المقبولة والمحققة: منحتان بقيمة ٤٠,٠٠٠ ريال (العنقري ٢٠ ألف + أبو زيد ٢٠ ألف) بنسبة تكلفة جمع تبرعات ٠.٠٪.", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "• الطلبات قيد الدراسة والمتابعة النشطة (١١ جهة):", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "  - صندوق دعم الجمعيات (مشروع جودة حياة - ٠١/٠٨/٢٠٢٦)\n  - مؤسسة محمد الشاوي (مبادرة عون للأمراض المزمنة - ٠٦/٠٨/٢٠٢٦)\n  - أوقاف الشيخ صالح الراجحي (دفء وغذاء ٢٠٢٧م - ١٢/٠٨/٢٠٢٦)\n  - مجلس الأوقاف الرائدة (تمت الموافقة على جودة الحياة)\n  - القطاع المصرفي: بنك الرياض وبنك البلاد (يوليو ٢٠٢٦)\n  - قطاع الشركات وخدمات الحجاج: ٦ شركات (الراجحي، ضيوف البيت، هوليدي إن، أبراج مكة، مشارق الماسية، الرفادة - ٠٧/٠٥/٢٠٢٦).", color=TEXT_DARK, size=10)
    
    add_rtl_p(doc, "• فرص منصة إحسان واعتذارات المانحين والتوصيات:", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "  - فرص إحسان تحت المعالجة: دعم المصاريف التشغيلية (تعديل التصنيف)، وجلسات غسيل الكلى (تحديث التقارير).\n  - اعتذارات لانتهاء الموازنات (٤ جهات): الماجد، الشاوي، طيبة، ومجموعة فنادق.\n  - اعتذارات لاشتراط الحوكمة والتخصص (٤ جهات): الضحيان، طلال (اشتراط الحوكمة)، الحمدان، والمهيدب (طفولة مبكرة).\n  - توصيات قسم البرامج: إعادة التقديم المبكر في Q1 2027، استكمال الحوكمة، ومتابعة فرص منصة إحسان (جود إحسان).", color=TEXT_DARK, size=10)

    # 7. Beneficiary Experiences & Verbatim Letters (Pages 40-42)
    add_rtl_heading(doc, "سابعاً: تجارب المستفيدين ورسائل الشكر والامتنان (ص ٤٠-٤٢)", level=1)
    add_rtl_p(doc, "١. تجربة المستفيدة سامية سليمان محمد (عملية استئصال كتلة بالصدر - ٦,٣٥٠ ريال بمستشفى المواساة):", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "«إلى أعضاء جمعية طبيبي الكرام، تعجز كلمات الثناء والشكر أن تفيكم حقكم لما تقدمونه من جهود جليلة في خدمة المجتمع، فأنتم نموذج يُحتذى به في البذل والعطاء... شكراً لعطائكم السخي الذي يضيء حياة الكثيرين ويصنع الأمل، فكل مساهمة منكم هي بذرة خير تُثمر بسمة في قلب محتاج.»", color=TEXT_DARK, size=9.5)
    
    add_rtl_p(doc, "٢. تجربة المستفيدة كندفة محمد عتبة (تنويم ورعاية بمدينة الملك سلمان الطبية - ٧,٠٠٠ ريال من ٩,١٨٦ ريال):", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "«من أعماق قلبي، أتقدم بخالص الشكر والامتنان لـ جمعية طبيبي. كنتم سبباً في تفريج كربتي في أصعب وقت... وأخص بالشكر الأستاذة غدير والأستاذ وائل وجميع الشباب العاملين في الجمعية، والله ما قصروا معنا أبداً وكانوا مثالاً للأخلاق والرحمة وحسن التعامل حتى اطمأن قلبي على والدتي.»", color=TEXT_DARK, size=9.5)
    doc.add_page_break()

    # 8. Strategic Aspirations, Advisory Team & 8 Initiatives (Pages 36-39)
    add_rtl_heading(doc, "ثامناً: التطلعات ومقترح الفريق الاستشاري والمبادرات الـ (٨) (ص ٣٦-٣٩)", level=1)
    add_rtl_p(doc, "• مقترح الاستعانة بفريق استشاري خارجي متخصص (بتكلفة ٥,٠٠٠ إلى ٧,٠٠٠ ريال شهرياً لمدة ٣ أشهر):", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "  - مبررات المقترح الـ (٦): تعدد التخصصات، الوفر المالي مقارنة بالتوظيف المباشر، استثمار المخرجات كأصول دائمة، رفع فرص قبول المنح الكبرى، نقل المعرفة للكادر، والرقابة التامة لمجلس الإدارة.\n  - خطة التنفيذ الثلاثية: المرحلة ١ (الحوكمة ونوى والقوائم) | المرحلة ٢ (إغلاق المشاريع والموقع وقاعدة المانحين) | المرحلة ٣ (الحقائب التعريفية والعروض والتقديم).\n  - ضمانات الرقابة الـ (٦) لمجلس الإدارة: اعتماد النطاق مسبقاً، تحديد مسؤول اتصال، تقارير دورية، مراجعة المدير التنفيذي، ربط الدفعات بالمخرجات، والسرية التامة وتسليم الملفات المفتوحة.", color=TEXT_DARK, size=10)

    add_rtl_p(doc, "• حزمة المبادرات والمقترحات التطويرية الـ (٨) المعتمدة للنصف الثاني:", bold=True, color=PRIMARY, size=10.5)
    inits = [
        "١. تنويع الشراكات ومصادر الدخل: التوسع في الأوقاف وشراكات القطاع الخاص لتحقيق الاستدامة المالية.",
        "٢. تعزيز الظهور الإعلامي: حضور أوسع بالفعاليات والمنصات الرقمية لزيادة الوعي بالجمعية.",
        "٣. استقطاب الكفاءات: شغل الشواغر في الإعلام، الموارد البشرية، وتنمية الموارد المالية.",
        "٤. تنفيذ برنامج «بطاقة طبيبي»: إصدار بطاقة مزايا وخصومات لدى المستشفيات والمختبرات والمتاجر الصحية.",
        "٥. تعديل لائحة صرف المساعدات: إدخال استثناءات إنسانية مرنة لرفع معدل قبول الحالات وصرف موازنة العلاج.",
        "٦. الاستعانة بالفريق الاستشاري: استكمال متطلبات الحوكمة، القوائم المالية، ومنصة نوى.",
        "٧. اعتماد مصفوفة الصلاحيات: لتعزيز سرعة ومرونة اتخاذ القرار ومواكبة متطلبات العمل.",
        "٨. تفعيل دور المجلس والعمومية: مساهمة الأعضاء الفاعلة في فتح مسارات الدعم المالي والمانحين."
    ]
    for init_item in inits:
        add_rtl_p(doc, init_item, color=TEXT_DARK, size=10)

    # 9. Supervisory Entities, Donors & Closing (Pages 43-47)
    add_rtl_heading(doc, "تاسعاً: الجهات الإشرافية والمانحون وخاتمة التقرير (ص ٤٣-٤٧)", level=1)
    add_rtl_p(doc, "• الجهات الإشرافية والحكومية: المركز الوطني لتنمية القطاع غير الربحي، وزارة الصحة، إمارة منطقة المدينة المنورة، ووزارة الموارد البشرية والتنمية الاجتماعية.", color=TEXT_DARK, size=10)
    add_rtl_p(doc, "• المنصات المعتمدة: منصة تبرع، منصة إحسان، منصة شفاء، والمتجر الإلكتروني الرسمي للجمعية.", color=TEXT_DARK, size=10)
    add_rtl_p(doc, "• أبرز الأوقاف والجهات المانحة: وقف الشيخ نغيمش الأحمدي، شركة طابة المطورة، وقف الشيخ عبدالقادر شيبة الحمد، مؤسسة سعيد مكي، وقف عبدالرحيم عبدالرزاق، ووقف عبدالعزيز أبو زيد.", color=TEXT_DARK, size=10)
    
    add_rtl_p(doc, "• كلمة الخاتمة الرسمية للإدارة التنفيذية:", bold=True, color=PRIMARY, size=10.5)
    add_rtl_p(doc, "«ختاماً، نؤكد في الإدارة التنفيذية أن هذه المنجزات لم تكن لتتحقق لولا فضل الله أولاً، ثم الدعم والتوجيه المستمر من قبل مجلس الإدارة الموقر، وتضافر جهود فريق العمل بجمعية طبيبي. إننا ننظر إلى هذا العام كبداية قوية حافلة بالعطاء، ونجدد التزامنا بمواصلة العمل التطويري وابتكار حلول مستدامة لتنمية الموارد وتوسيع نطاق خدماتنا الطبية لتصل إلى كل محتاج. نسأل الله أن يبارك في الجهود وأن يسدد الخطى لما فيه الخير والنفع لمجتمعنا ووطننا.»", color=TEXT_DARK, size=10)

    doc.save(docx_path)
    print(f"Saved Comprehensive Word Document: {docx_path}")

# Run word generator
word_out_v2 = os.path.join(v2_dir, "تقرير_جمعية_طبيبي_النصف_سنوي_٢٠٢٦_النسخة_التنفيذية.docx")
build_comprehensive_word_doc(word_out_v2)
shutil.copy2(word_out_v2, os.path.join(v1_dir, "تقرير_جمعية_طبيبي_النصف_سنوي_٢٠٢٦_النسخة_التنفيذية.docx"))

# Run HTML generator
os.system(f'py -3 "{os.path.join(base_dir, "generate_v2_dashboard.py")}"')

print("Word Document and HTML Dashboard fully synchronized with all pages 30 to 47 data!")
